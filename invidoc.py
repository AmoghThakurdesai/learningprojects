# invitation maker-put the names in the word doc and makes the invitations for you

import docx
doc=docx.Document('Invitations.docx')
guestlist=[]
file=open('D:\\PythonFiles\\guests.txt')
for name in file.readlines():
    guestlist.append(name)
for i in range(0,len(guestlist)):
    para1=doc.add_paragraph("It would be a pleasure to have the company of")
    para1.style='Invigeneral'
    para2=doc.add_paragraph(guestlist[i])
    para2.style='Inviname'
    para3=doc.add_paragraph("at 11010 Memory Lane on the Evening of")
    para3.style='Invigeneral'
    para4=doc.add_paragraph('April 1st')
    para4.style='Inviname'
    para5=doc.add_paragraph("at 7 o'clock")
    para5.style='Invigeneral'
    para5.runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.save('Invitations.docx')

